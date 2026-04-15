"""
MOM & Weekly Log – strict 7-day weeks, ~1 page each.
Project: 21 January – 7 April 2026  →  11 weeks of 7 days.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x1F, 0x35, 0x64)
GOLD      = RGBColor(0xC0, 0x8C, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)

TEAM = [
    "Cristine Valentina",
    "Habil Baihaki Ramadhan",
    "Johana Veronica Setiawan",
    "Nasywa Chonifahtun Fiqrihiyah",
    "Nisrina Izza Nur Aisyah",
    "Shanty",
]

# ── DATA: 11 strict 7-day weeks ───────────────────────────────────────────────
WEEKS = [
    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 1",
        "period" : "21 – 27 January 2026",
        "focus"  : "Team Formation & Initial Brainstorming (IDX Stock Analyzer)",
        "activities": [
            ("21 Jan", "All",            "Kick-off meeting via Google Meet. Team of 6 formed, roles tentatively assigned. Course brief reviewed. Proposed project: IDX Stock Analyzer (real-time Indonesian stock dashboard)."),
            ("22 Jan", "Habil, Nasywa",  "API research for IDX stocks: Yahoo Finance (limited), IDX portal (static CSV), Alpha Vantage (no IDX on free tier), Investing.com (scraping prohibited). No viable free real-time source found."),
            ("23 Jan", "Johana, Cristine","Drew UI wireframe sketches for IDX dashboard: Market Overview, Stock Detail, and Watchlist views."),
            ("24 Jan", "Shanty, Nisrina","Evaluated charting libraries (Recharts, Chart.js, ApexCharts)—Recharts shortlisted. Compared Supabase vs Firebase for auth."),
            ("27 Jan", "All",            "Compiled research. Blockers confirmed: no free real-time IDX API, CORS restriction. Agreed to consult lecturer before committing to a direction."),
        ],
        "mom_date"  : "21 January 2026",
        "mom_time"  : "19.00 – 21.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Team introductions and role assignment",
            "Review project brief and rubric requirements",
            "Idea generation: candidate project topics",
            "Assign Week 1 research tasks",
        ],
        "decisions" : [
            "Project proposed: IDX Stock Analyzer (pending lecturer approval).",
            "Tech stack tentatively agreed: React + TypeScript + Vite.",
            "Weekly meetings every Monday at 19.00 WIB via Google Meet.",
        ],
        "action_items": [
            ("Habil, Nasywa",    "Research free real-time APIs for IDX stocks",             "27 Jan"),
            ("Johana, Cristine", "Create UI wireframe sketches for IDX dashboard",           "27 Jan"),
            ("Shanty",           "Shortlist charting library (Recharts vs alternatives)",    "27 Jan"),
            ("Nisrina",          "Compare Supabase vs Firebase: features and free tier",     "27 Jan"),
            ("All",              "Consolidate findings into shared Google Doc",              "27 Jan"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 2",
        "period" : "28 January – 3 February 2026",
        "focus"  : "Lecturer Consultation & Pivot Decision: IDX → Gold Analysis Platform",
        "activities": [
            ("28 Jan", "All",           "Attended AI-1 class. Presented IDX Analyzer concept to Mr. Hendra Jayanto. Lecturer flagged the lack of free real-time IDX APIs as a high risk and recommended pivoting to gold (XAU/USD)—globally supported, abundant free APIs, strong AI potential."),
            ("29 Jan", "Habil, Nasywa", "Rapid gold API research: Stooq.com = free XAU/USD CSV endpoint; Yahoo Finance = free historical OHLC. CORS solution: Vercel Serverless Functions proxy."),
            ("30 Jan", "All",           "Emergency meeting. Unanimous vote to PIVOT to Gold Analysis Platform. Scope defined: live price, chart, AI prediction, Telegram alerts, paper trading, bilingual UI."),
            ("3 Feb",  "All",           "Updated project brief document to reflect the pivot. New roles confirmed: Habil + Nasywa (backend/API), Johana + Cristine (UI), Shanty (charts), Nisrina (auth/DB)."),
        ],
        "mom_date"  : "30 January 2026",
        "mom_time"  : "20.00 – 22.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Recap lecturer feedback from 28 January class",
            "Formal vote: continue IDX vs. pivot to Gold platform",
            "Define new project scope",
            "Reassign component ownership",
        ],
        "decisions" : [
            "PIVOT CONFIRMED (unanimous): Gold Analysis Platform adopted as official project direction.",
            "Nine feature tabs agreed: AI Prediction, Simulator, Telegram, Technical, Fundamental, News, Calendar, Correlation, Experts.",
            "Vercel Serverless Functions = CORS solution for all external API calls.",
        ],
        "action_items": [
            ("Habil",            "Set up GitHub repo (Vite + React + TS) and invite all members",     "4 Feb"),
            ("Johana, Cristine", "Redesign UI wireframes for Gold platform (dark mode, 9-tab layout)", "6 Feb"),
            ("Nasywa",           "Document Stooq and Yahoo Finance API endpoints and field mappings",  "3 Feb"),
            ("Nisrina",          "Prepare Supabase project setup plan and SQL schema draft",           "3 Feb"),
            ("Shanty",           "Confirm Recharts as charting library; plan component structure",     "3 Feb"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 3",
        "period" : "4 – 10 February 2026",
        "focus"  : "Repository Setup, Supabase Initialization & Tech Stack Finalization",
        "activities": [
            ("4 Feb",  "Habil",   "Initialized GitHub repository: Vite + React 18 + TypeScript + Tailwind CSS. Configured ESLint, Prettier, main/develop branches. All members invited as collaborators."),
            ("6 Feb",  "Nisrina", "Created Supabase project. Wrote SQL schema for users table (telegram_bot_token, telegram_chat_id, preferred_language, initial_balance). Row-Level Security (RLS) configured. Credentials shared with team."),
            ("7 Feb",  "Habil",   "Created Vercel project and linked to GitHub repo. CI/CD pipeline confirmed—pushes to main trigger automatic Vercel builds."),
            ("8 Feb",  "Shanty",  "Finalized official Technology Stack document: React 18 + TypeScript, Vite, Recharts, Supabase, Vercel Serverless Functions."),
            ("10 Feb", "All",     "All members cloned repo and confirmed local dev server running. Component ownership formally documented and agreed."),
        ],
        "mom_date"  : "10 February 2026",
        "mom_time"  : "19.00 – 20.30 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Confirm all members have working local dev environment",
            "Review Supabase schema and Vercel CI/CD setup",
            "Assign Week 4 development tasks (layout and proxy)",
        ],
        "decisions" : [
            "Technology stack finalized and locked: React 18 + TypeScript + Vite + Recharts + Supabase + Vercel.",
            "GitHub Flow adopted: feature branches → PR → code review → merge to main.",
            "Active development begins from Week 4.",
        ],
        "action_items": [
            ("Johana, Cristine", "Build main application layout: nav bar, sidebar, tab routing",  "14 Feb"),
            ("Habil",            "Implement /api/stooq.ts serverless proxy; test CORS bypass",    "14 Feb"),
            ("Nasywa",           "Build useGoldPrices hook: fetch, parse CSV, expose as state",   "14 Feb"),
            ("Shanty",           "Build CandlestickChart component shell with mock OHLC data",    "14 Feb"),
            ("Nisrina",          "Set up Supabase auth client in the React project",               "14 Feb"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 4",
        "period" : "11 – 17 February 2026",
        "focus"  : "App Layout, Stooq Proxy & Live Price Feed",
        "activities": [
            ("11 Feb", "Johana, Cristine","Built main app layout: top nav bar (tabs, language & theme toggle), sidebar (XAU/XAG price cards), and main content area with tab routing. Dark mode palette applied (gold accent #C9A800)."),
            ("12 Feb", "Habil",          "Implemented /api/stooq.ts Vercel Serverless proxy. CORS bypass confirmed—browser receives live XAU/USD CSV without errors."),
            ("14 Feb", "Nasywa",         "Built useGoldPrices hook: calls /api/stooq every 15 seconds, parses CSV (Open, High, Low, Close), exposes price as React state. Live price updating on screen for the first time."),
            ("15 Feb", "Habil",          "Added 1-second tick simulation (±0.3% micro-movement) between Stooq fetches. Price display now appears continuously live, mimicking a real trading terminal."),
            ("17 Feb", "Shanty",         "Built CandlestickChart component using Recharts with mock OHLC data. Custom candlestick shapes (green/red). Timeframe buttons (1D, 1W, 1M, 3M) and instrument toggle (XAU/XAG) implemented."),
        ],
        "mom_date"  : "17 February 2026",
        "mom_time"  : "19.30 – 20.30 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Demo live price feed on screen (Nasywa)",
            "Review tick simulation effect (Habil)",
            "Plan Yahoo Finance OHLC integration for Week 5",
        ],
        "decisions" : [
            "Live price feed via Stooq proxy confirmed working in development environment.",
            "1-second tick simulation approved—±0.3% variation agreed as optimal value.",
            "Next priority: replace mock OHLC with real Yahoo Finance data.",
        ],
        "action_items": [
            ("Habil",         "Implement /api/yahoo.ts proxy for historical OHLC data",                  "21 Feb"),
            ("Shanty, Nasywa","Integrate Yahoo Finance OHLC data into CandlestickChart",                 "21 Feb"),
            ("Nasywa",        "Add XAG/USD derivation using gold/silver ratio to the hook",              "21 Feb"),
            ("Johana, Cristine","Polish UI layout: spacing, color, responsive breakpoints",              "24 Feb"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 5",
        "period" : "18 – 24 February 2026",
        "focus"  : "Yahoo Finance Integration, Live Chart & Sprint Review",
        "activities": [
            ("19 Feb", "Habil",          "Implemented /api/yahoo.ts Vercel proxy for historical OHLC from Yahoo Finance. Retry logic with exponential backoff + 5-min cache added to handle HTTP 429 rate-limit responses."),
            ("21 Feb", "Shanty, Nasywa", "Integrated live OHLC data into CandlestickChart—replaced mock data. Most recent candle updates in real time with each price tick. Timeframe switching re-fetches and re-renders correctly."),
            ("22 Feb", "Nasywa",         "Added XAG/USD derivation inside useGoldPrices: XAG = XAU ÷ goldSilverRatio. Silver price now updates live with no additional API call."),
            ("24 Feb", "All",            "Sprint review: full team demoed the live price feed + real OHLC candlestick chart. Core data pipeline approved as production-ready."),
        ],
        "mom_date"  : "24 February 2026",
        "mom_time"  : "19.30 – 21.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Sprint review: demo live chart with real OHLC data",
            "Confirm XAG/USD derivation approach",
            "Plan Week 6: technical indicator overlays",
        ],
        "decisions" : [
            "Core data pipeline approved as production-ready. Architecture locked.",
            "XAG/USD derived from live XAU price—no separate API call needed.",
            "Week 6 focus: implement all 6 indicator overlays on the chart.",
        ],
        "action_items": [
            ("Shanty",   "Add MA20 and MA50 overlay calculations and chart toggles",          "28 Feb"),
            ("Shanty",   "Add EMA12, EMA26, Bollinger Bands, and Fibonacci overlays",         "3 Mar"),
            ("Nasywa",   "Add RSI (14), MACD (12,26,9), ADX to Technical Analysis panel",    "3 Mar"),
            ("Cristine", "Begin design of AI Prediction engine algorithm",                    "3 Mar"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 6",
        "period" : "25 February – 3 March 2026",
        "focus"  : "Technical Indicator Overlays",
        "activities": [
            ("27 Feb", "Shanty",  "Implemented MA20 and MA50 overlays on CandlestickChart. Values computed client-side from OHLC data. Toggle buttons added to chart toolbar."),
            ("3 Mar",  "Shanty",  "Added EMA12, EMA26, Bollinger Bands (SMA ± 2σ), and Fibonacci Retracement (23.6%, 38.2%, 50%, 61.8%). All 6 overlays independently toggleable. Validated against TradingView reference."),
            ("3 Mar",  "Nasywa",  "Added RSI (14-period), MACD (12,26,9), and ADX calculations to Technical Analysis panel. Gauge and histogram displays added alongside the chart."),
        ],
        "mom_date"  : "3 March 2026",
        "mom_time"  : "19.00 – 20.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Demo all 6 indicator overlays (Shanty)",
            "Review RSI/MACD/ADX panel output (Nasywa)",
            "Plan Week 7: AI Prediction engine and News panel",
        ],
        "decisions" : [
            "All 6 overlays approved. Chart module considered feature-complete.",
            "Cristine to lead AI Prediction engine: rule-based synthesis of RSI, MACD, EMA cross, Bollinger, news sentiment.",
        ],
        "action_items": [
            ("Cristine",  "Build AI Prediction engine (Bullish/Bearish probability, per-indicator reasoning)", "10 Mar"),
            ("Cristine",  "Build dynamic News & Sentiment panel (price-aware headline generator)",             "10 Mar"),
            ("Habil",     "Research Telegram Bot API; design /api/telegram.ts proxy",                         "10 Mar"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 7",
        "period" : "4 – 10 March 2026",
        "focus"  : "AI Prediction Engine & News Sentiment Panel",
        "activities": [
            ("8 Mar",  "Cristine", "Completed AI Prediction engine: synthesizes RSI, MACD, EMA cross, Bollinger Band position, news sentiment, and DXY correlation into a Bullish/Bearish probability (0–100%). Output includes confidence level, risk/reward ratio, and plain-English per-indicator reasoning."),
            ("10 Mar", "Cristine", "Built dynamic News & Sentiment panel: headlines generated from price-direction templates (rising price → bullish headlines, falling → bearish). Sentiment score derived from AI prediction output."),
            ("10 Mar", "Habil",    "Researched Telegram Bot API. Designed /api/telegram.ts proxy architecture: accepts slug param, forwards to Telegram API server-side to bypass CORS. Ready for implementation in Week 8."),
        ],
        "mom_date"  : "10 March 2026",
        "mom_time"  : "19.30 – 20.30 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Demo AI Prediction output (Cristine)",
            "Demo News & Sentiment panel (Cristine)",
            "Plan Week 8: Telegram integration and Supabase auth",
        ],
        "decisions" : [
            "AI Prediction engine approved. Algorithm output considered balanced and contextually appropriate.",
            "News panel approved—dynamic generation accepted as alternative to paid news API.",
            "Week 8: Habil + Johana to implement Telegram proxy and settings UI; Nisrina to implement Supabase auth.",
        ],
        "action_items": [
            ("Habil, Johana", "Implement /api/telegram.ts proxy and Telegram Settings UI",        "13 Mar"),
            ("Nisrina",       "Implement Supabase email/password auth and dual persistence",       "13 Mar"),
            ("All",           "Sprint review session at end of Week 8",                            "13 Mar"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 8",
        "period" : "11 – 17 March 2026",
        "focus"  : "Telegram Integration, Supabase Auth & Bug Fix",
        "activities": [
            ("12 Mar", "Habil, Johana", "Implemented /api/telegram.ts Vercel Serverless proxy. Routes all Telegram Bot API calls server-side, bypassing CORS. Johana built Telegram Settings UI: bot token input, Chat ID detect button, connection status indicator."),
            ("12 Mar", "Nisrina",       "Implemented Supabase email/password auth. User profile (telegram_bot_token, chat_id, language, balance) stored in DB. Dual persistence: localStorage (offline-safe) + Supabase (cross-device sync)."),
            ("13 Mar", "All",           "Sprint review: all features demoed (~60% of scope complete). Telegram Chat ID auto-detection bug reported—returns empty array when user hasn't messaged the bot yet."),
            ("17 Mar", "Habil, Johana", "Fixed Telegram Chat ID bug: added onboarding instruction ('Send any message to your bot first') + polling retry with 5-second delay. Bug resolved and verified."),
        ],
        "mom_date"  : "13 March 2026",
        "mom_time"  : "20.00 – 21.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Sprint review: demo Telegram integration and Supabase auth",
            "Report Telegram Chat ID bug and plan fix",
            "Set Week 9 priorities",
        ],
        "decisions" : [
            "Telegram proxy and auth approved. ~60% of total scope complete.",
            "Chat ID bug fix: assigned to Habil + Johana. Due 17 March.",
            "Week 9 priority: Price Alerts, Correlated Assets, Economic Calendar.",
        ],
        "action_items": [
            ("Habil, Johana",  "Fix Telegram Chat ID bug (onboarding step + retry polling)",              "17 Mar"),
            ("Johana, Nasywa", "Build Price Alert engine (threshold monitoring + Telegram notification)", "20 Mar"),
            ("Nasywa, Shanty", "Build Correlated Assets panel (DXY, US10Y, Oil, BTC derived values)",    "21 Mar"),
            ("Shanty",         "Build Economic Calendar (high-impact events, dynamically filtered)",       "21 Mar"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 9",
        "period" : "18 – 24 March 2026",
        "focus"  : "Price Alerts, Correlated Assets & Economic Calendar",
        "activities": [
            ("19 Mar", "Johana, Nasywa", "Implemented Price Alert engine: monitors live price every tick against user-defined thresholds; fires formatted Telegram notification on breach; auto-disables after trigger to prevent repeated alerts."),
            ("21 Mar", "Nasywa",         "Completed Correlated Assets panel: DXY, US Treasury Yield (10Y), Crude Oil, Bitcoin values derived dynamically from live XAU price using correlation formulas. Each card shows value, correlation coefficient, and Bullish/Bearish signal."),
            ("21 Mar", "Shanty",         "Completed Economic Calendar: high-impact macroeconomic events (FOMC, US CPI, NFP, China PMI) shown for current and next week, filtered dynamically by system date."),
            ("24 Mar", "All",            "Mid-sprint check. All new features confirmed working. Remaining: Expert Analysis panel, Trading Simulator, i18n, final QA, and Vercel deployment."),
        ],
        "mom_date"  : "24 March 2026",
        "mom_time"  : "19.30 – 20.30 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Demo Price Alerts, Correlated Assets, Economic Calendar",
            "Confirm deployment timeline",
            "Plan Week 10: Expert Analysis, i18n, deployment",
        ],
        "decisions" : [
            "Price Alerts, Correlated Assets, and Economic Calendar approved.",
            "Production deployment target: goldaiprediction.vercel.app by 28 March.",
            "No new features after 1 April—full QA focus from 2 April.",
        ],
        "action_items": [
            ("Habil",    "Configure Vercel environment variables; deploy to goldaiprediction.vercel.app", "28 Mar"),
            ("Cristine", "Build Expert Analysis panel (5 analysts, dynamic price targets)",               "28 Mar"),
            ("Nisrina",  "Implement bilingual i18n system (English / Indonesian) using react-i18next",    "1 Apr"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 10",
        "period" : "25 – 31 March 2026",
        "focus"  : "Expert Analysis, i18n & Production Deployment",
        "activities": [
            ("28 Mar", "Cristine", "Completed Expert Analysis panel: 5 analyst profiles with signal ratings (Strong Buy–Sell), price targets computed as % offsets from live price, historical accuracy %, and analysis rationale."),
            ("28 Mar", "Habil",    "Deployed full application to Vercel production. Environment variables (Supabase URL, anon key) configured in Vercel dashboard. goldaiprediction.vercel.app confirmed live."),
            ("29 Mar", "Nisrina",  "Completed bilingual i18n system with react-i18next. Language toggle in nav bar switches entire UI (all 9 tabs) instantly between English and Indonesian. Preference saved to localStorage and Supabase."),
            ("31 Mar", "All",      "End-of-week check: deployment live and stable. Expert Analysis and i18n confirmed working on production URL. One remaining major feature: Paper Trading Simulator."),
        ],
        "mom_date"  : "31 March 2026",
        "mom_time"  : "20.00 – 21.00 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Confirm goldaiprediction.vercel.app is live and stable",
            "Review Expert Analysis and i18n on production",
            "Plan Week 11: Trading Simulator + full integration QA",
        ],
        "decisions" : [
            "Production deployment confirmed stable. CI/CD working (auto-build on push to main).",
            "Week 11 is the final sprint: Trading Simulator due 1 April, full QA 2–5 April.",
            "Presentation day: 7 April. No new features after 1 April.",
        ],
        "action_items": [
            ("Habil, Johana", "Complete Paper Trading Simulator (Long/Short, P&L, TP/SL, Telegram on close)", "1 Apr"),
            ("All",           "Full integration QA: test every feature, not just your own",                    "5 Apr"),
            ("All",           "Prepare individual presentation slides and rehearse speaking segments",          "6 Apr"),
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "label"  : "Week 11",
        "period" : "1 – 7 April 2026",
        "focus"  : "Trading Simulator, Final QA & Presentation",
        "activities": [
            ("1 Apr",   "Habil, Johana", "Completed Paper Trading Simulator: virtual balance, Long/Short positions at live price, real-time P&L per tick (P&L = Δprice × lotSize), Take Profit/Stop Loss auto-close with Telegram notification, manual close supported."),
            ("1 Apr",   "All",           "Feature freeze. No new code commits after this date."),
            ("2–4 Apr", "All",           "Full integration QA. 11 bugs found and resolved: 5 UI alignment, 3 logic errors, 1 Telegram token URI-encoding, 1 race condition in tick simulation (NaN flash), 1 i18n hydration on initial load. All responsive breakpoints passed."),
            ("5 Apr",   "All",           "Second QA pass: all 11 bugs confirmed resolved. No new critical issues. Documentation finalized (Software Engineer Report, MOM Log, FAQ)."),
            ("6 Apr",   "All",           "Final check on goldaiprediction.vercel.app. Presentation roles rehearsed. Habil: architecture; Cristine: AI & features; Johana: live demo; Nasywa: data pipeline; Shanty: charts & indicators; Nisrina: auth, i18n & Q&A."),
            ("7 Apr",   "All",           "Final presentation delivered to Mr. Hendra Jayanto. Live demo of all 9 tabs on goldaiprediction.vercel.app. Positive feedback received. Project completed."),
        ],
        "mom_date"  : "6 April 2026",
        "mom_time"  : "20.00 – 21.30 WIB",
        "mom_loc"   : "Online – Google Meet",
        "agenda"    : [
            "Final QA sign-off: confirm all 11 bugs resolved",
            "Live check of goldaiprediction.vercel.app",
            "Presentation roles and demo script confirmed",
            "Team retrospective",
        ],
        "decisions" : [
            "Platform declared PRODUCTION-READY. All features confirmed working.",
            "No further code changes before presentation.",
            "Presentation roles and order formally agreed by all members.",
            "Future work logged: server-side Telegram alerts, backtesting engine, mobile app.",
        ],
        "action_items": [
            ("All",     "Rehearse individual presentation segments independently",        "7 Apr"),
            ("Johana",  "Prepare and rehearse live demo script for all 9 platform tabs",  "7 Apr"),
            ("Habil",   "Prepare architecture diagram and data-flow slides",               "7 Apr"),
            ("Nisrina", "Finalize FAQ document for anticipated Q&A",                       "7 Apr"),
        ],
    },
]

# ── HELPERS ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph(style='Heading 1'); p.clear()
    r = p.add_run(text.upper())
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F3564')
    pPr.append(shd)

def h2(doc, text):
    p = doc.add_paragraph(style='Heading 2'); p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)

def lv(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f"{label}: "); r1.font.name = 'Times New Roman'; r1.font.size = Pt(9.5); r1.font.bold = True; r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(value);        r2.font.name = 'Times New Roman'; r2.font.size = Pt(9.5); r2.font.color.rgb = BLACK

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(9.5); r.font.color.rgb = BLACK

def make_table(doc, headers, rows, col_widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ''; set_cell_bg(c, '1F3564')
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        tr = tbl.add_row(); fill = 'EEF2FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = tr.cells[ci]; c.text = ''; set_cell_bg(c, fill)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.font.color.rgb = BLACK
    for row in tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

def section_bold(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = DARK_BLUE

def pagebreak(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ── BUILD ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width  = Cm(21.0)
    sec.left_margin = Cm(2.5);  sec.right_margin = Cm(2.5)
    sec.top_margin  = Cm(2.0);  sec.bottom_margin = Cm(2.0)

    # ── COVER ────────────────────────────────────────────────────────────────
    def cp(text, sz=11, bold=False, color=None, gap=4):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(gap)
        if text:
            rn = p.add_run(text); rn.font.name = 'Times New Roman'; rn.font.size = Pt(sz); rn.font.bold = bold
            if color: rn.font.color.rgb = color

    cp("SOFTWARE ENGINEERING DOCUMENTATION", 14, True, DARK_BLUE, 3)
    cp("Minutes of Meeting (MOM) & Weekly Activity Log", 13, True, GOLD, 3)
    cp("Gold Analysis Platform", 12, True, gap=2)
    cp("Real-Time Gold & Market Intelligence System – XAU/USD & XAG/USD", 10, gap=14)
    cp("Lecturer:", 10, True, gap=1); cp("Hendra Jayanto", 10, gap=10)
    cp("Class:", 10, True, gap=1);    cp("Artificial Intelligence 1", 10, gap=10)
    cp("Created by:", 10, True, gap=4)
    for name in TEAM: cp(name, 10, gap=1)
    cp("", gap=14)
    cp("INFORMATICS – FACULTY OF COMPUTER SCIENCES", 10, True, gap=1)
    cp("PRESIDENT UNIVERSITY", 10, True, gap=1); cp("2025/2", 10, gap=0)
    pagebreak(doc)

    # ── OVERVIEW TABLE ───────────────────────────────────────────────────────
    h1(doc, "Project Overview")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    make_table(doc,
        ["Week", "Period", "Main Focus"],
        [(w["label"], w["period"], w["focus"]) for w in WEEKS],
        col_widths=[1.5, 4.0, 9.0],
    )
    pagebreak(doc)

    # ── PER-WEEK PAGES ───────────────────────────────────────────────────────
    for i, wk in enumerate(WEEKS):
        h1(doc, f"{wk['label']}  |  {wk['period']}")
        lv(doc, "Focus", wk["focus"])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        h2(doc, "Activity Log")
        make_table(doc,
            ["Date", "Member(s)", "Activity"],
            wk["activities"],
            col_widths=[1.5, 3.2, 9.8],
        )

        h2(doc, f"Minutes of Meeting  —  {wk['mom_date']}")
        lv(doc, "Time",     wk["mom_time"])
        lv(doc, "Location", wk["mom_loc"])
        lv(doc, "Attendees", ", ".join(TEAM))
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

        section_bold(doc, "Agenda")
        for item in wk["agenda"]:
            bullet(doc, item)

        section_bold(doc, "Decisions Made")
        for d in wk["decisions"]:
            bullet(doc, d)

        section_bold(doc, "Action Items")
        make_table(doc,
            ["Assigned To", "Task", "Due"],
            wk["action_items"],
            col_widths=[3.5, 9.5, 1.5],
        )

        if i < len(WEEKS) - 1:
            pagebreak(doc)

    out = "MOM & Weekly Log - Gold Analysis Platform.docx"
    doc.save(out)
    print(f"Saved: {out}  ({len(WEEKS)} weeks)")

if __name__ == "__main__":
    build()
